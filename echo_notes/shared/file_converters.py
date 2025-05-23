"""
File converters for different file formats.
This module provides functions to convert different file formats to text
that can be processed by the LLM.
"""

from pathlib import Path
import os

def txt_to_text(file_path: Path) -> str:
    """Convert a .txt file to text"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def docx_to_text(file_path: Path) -> str:
    """
    Convert a .docx file to text.
    Requires python-docx package.
    """
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx package is required to process .docx files. "
            "Please install it with: pip install python-docx"
        )
    
    doc = docx.Document(file_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)

def text_to_docx(file_path: Path, content: str) -> None:
    """
    Write text content to a .docx file.
    Requires python-docx package.
    """
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx package is required to process .docx files. "
            "Please install it with: pip install python-docx"
        )
    
    # Create a new document
    doc = docx.Document()
    
    # Add paragraphs for each line in the content
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    
    # Save the document
    doc.save(file_path)

def get_converter_for_file(file_path: Path):
    """Get the appropriate converter function for a file based on its extension"""
    extension = file_path.suffix.lower()
    
    converters = {
        '.txt': txt_to_text,
        '.docx': docx_to_text,
        '.md': txt_to_text,  # Markdown files can be read as text
    }
    
    return converters.get(extension)

def get_writer_for_file(file_path: Path):
    """Get the appropriate writer function for a file based on its extension"""
    extension = file_path.suffix.lower()
    
    writers = {
        '.txt': lambda path, content: open(path, 'w', encoding='utf-8').write(content),
        '.docx': text_to_docx,
        '.md': lambda path, content: open(path, 'w', encoding='utf-8').write(content),
    }
    
    return writers.get(extension)